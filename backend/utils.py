
# This file contains all my helper functions that do the actual work.
import os
import asyncio, httpx
import logging
import docx
import html
import io
from fastapi import UploadFile, HTTPException
from pypdf import PdfReader
import assemblyai as aai
from datetime import datetime
import config
import google.generativeai as genai
from langdetect import detect, LangDetectException

logger = logging.getLogger(__name__)


DEFAULT_ROLES = {    
    "general overview", "No Selection", "general", "normal", "full review", "default", "standard", "all", "overall", "unfiltered", "everyone", "comprehensive"
}

async def read_text_from_file(file: UploadFile):
    filename = file.filename.lower()
    try:
        file_content = await file.read()
        file_stream = io.BytesIO(file_content)

        if filename.endswith(".pdf"):
            reader = PdfReader(file_stream)
            text_parts = [page.extract_text() for page in reader.pages if page.extract_text()]
            return "".join(text_parts).strip()
        elif filename.endswith(".docx"):
            doc = docx.Document(file_stream)
            full_text_list = [para.text for para in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text_list.append(cell.text)
            return "\n".join(full_text_list).strip()
        elif filename.endswith(".txt"):
            return file_content.decode("utf-8").strip()
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: '{filename}'. Please upload a PDF, DOCX, or TXT file.",
            )
    except Exception as e:
        logger.error(f"Error reading file {file.filename}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process file: {e}")



# this transcribes an audio file using the AssemblyAI API and formats the output with speaker labels if available.
async def transcribe_with_assemblyai(file: UploadFile, language: str):
    
    try:
        
        aai.settings.api_key = os.getenv("ASSEMBLYAI_API_KEY")
        if not aai.settings.api_key:
            raise HTTPException(status_code=500, detail="AssemblyAI API key not configured.")
        
        speaker_labels_lang = ['auto','en_us','en_uk','es','fr','de','it','ja','pt','ru','zh']

        config_params = {}
        if language in speaker_labels_lang:
            config_params["speaker_labels"] = True

        if language == "auto":
            config_params["language_detection"] = True
            logger.info("AssemblyAI transcription running with automatic language detection.")
        else:
            
            
            config_params["language_code"] = language
            logger.info(f"AssemblyAI transcription running with language detection, boosted for: {language}")

        config = aai.TranscriptionConfig(**config_params)
        transcriber = aai.Transcriber()

        
        transcript = await asyncio.to_thread(transcriber.transcribe, file.file, config)

        if transcript.status == aai.TranscriptStatus.error:
            logger.error(f"AssemblyAI transcription failed: {transcript.error}")
            raise HTTPException(status_code=500, detail=transcript.error)

 
        if config.speaker_labels and transcript.utterances:
            formatted_transcript = "\n".join(
                f"Speaker {utterance.speaker}: {utterance.text}"
                for utterance in transcript.utterances
            )
            return formatted_transcript
        else:
      
            return transcript.text

    except Exception as e:
        logger.error(f"AssemblyAI transcription failed for {file.filename}: {e}")
        raise HTTPException(status_code=500, detail=f"Transcription failed: {e}")



async def generate_gemini_content(prompt: str, model_name: str = "gemini-1.5-flash", is_json: bool = False):
    try:
        model = genai.GenerativeModel(model_name)
        response_type = "application/json" if is_json else "text/plain"
        generation_config = genai.types.GenerationConfig(response_mime_type=response_type)
        response = await model.generate_content_async(prompt, generation_config=generation_config)
        return response.text
    except Exception as e:
        logger.error(f"Gemini API call failed: {e}")
        raise HTTPException(status_code=500, detail=f"AI content generation failed: {e}")



async def role_summary(general_summary: str, role:str):
        
    if not role or role.strip().lower() in DEFAULT_ROLES:
        logger.info("No specific role selected. Returning general summary.")
        return general_summary

    logger.info(f"Refining summary for role: {role}")

    
    refine_prompt = f"""
            You are an expert at refining summaries for a specific person's role. Your task is to revise the provided General Summary to be specifically relevant for the person in the '{role}'.

            First, analyze the "General Summary to Refine" to determine if the '{role}' is mentioned or has any assigned tasks or direct relevance.

            ---
            IF THE ROLE IS MENTIONED or has relevant tasks/decisions:
            - The output MUST be a filtered summary.
            - Focus ONLY on the key decisions and action items that directly impact or are assigned to the '{role}'.
            - Omit all details not relevant to this specific role.
            - Maintain the original language and structure (headings, bullet points).
            - You MUST only use information present in the 'General Summary to Refine'.
            - Do not use markdown like ``, # or **.

            IF THE ROLE IS NOT MENTIONED and has no relevance:
            - You MUST add a note at the very top: "This summary is not specific to the '{role}' role, as it was not a focus of the discussion. Here is the general summary:"
            - You MUST return the original "General Summary to Refine" exactly as it is.
            - Do not add or assume any other information.
            - Do not use markdown like ``, # or **.
            ---

            General Summary to Refine:
            {general_summary}
            ---
    """

    refined_summary = await generate_gemini_content(refine_prompt)
    return refined_summary

#This checks if the summary language matches the original text language.
#If not, it translates the summary to the correct language.
async def correct_summary_language(original_text: str, summary_text: str):
    
    try:
       
        try:
            
            original_lang = detect(original_text[:1000])
        except LangDetectException:
            logger.warning("Could not detect language of original text. Skipping correction.")
            return summary_text
            
        try:
            summary_lang = detect(summary_text)
        except LangDetectException:
            logger.warning("Could not detect language of summary. Skipping correction.")
            return summary_text

        logger.info(f"Detected original language: {original_lang}, Detected summary language: {summary_lang}")

        
        if original_lang != summary_lang:
            logger.warning(f"Language mismatch detected. Translating summary from '{summary_lang}' to '{original_lang}'.")
            
            correction_prompt = f"""Translate the following text into the language with the ISO 639-1 code '{original_lang}'.
                Maintain the Structure of the text. Do not add any interpretation or assumptions. Provide ONLY the translated text.

                Text to translate:
                ---
                {summary_text}
                ---
                """
            
            corrected_summary = await generate_gemini_content(correction_prompt)
            return corrected_summary
        else:
            
            return summary_text

    except Exception as e:
        logger.error(f"An unexpected error occurred during language correction: {e}")
        # In case of any other error, it's safer to return the original summary
        return summary_text


async def send_email_via_brevo(recipient_email, subject, html_content, sender_name="AI Meeting Wizard"):
    brevo_api_key = os.getenv("BREVO_API_KEY")
    sender_email = os.getenv("SENDER_EMAIL")
    if not brevo_api_key or not sender_email:
        logger.error("Email service not configured. Missing BREVO_API_KEY or SENDER_EMAIL.")
        return False

    headers = {"api-key": brevo_api_key, "Content-Type": "application/json", "Accept": "application/json"}
    payload = {"sender": {"name": sender_name, "email": sender_email}, "to": [{"email": recipient_email}], "subject": subject, "htmlContent": html_content}
    
    async with httpx.AsyncClient() as client:
        try:
            response = await client.post("https://api.brevo.com/v3/smtp/email", json=payload, headers=headers, timeout=30.0)
            response.raise_for_status()
            logger.info(f"Email sent successfully to {recipient_email}")
            return True
        except httpx.HTTPStatusError as e:
            logger.error(f"Brevo API Error sending to {recipient_email}: {e.response.text}")
            return False
        except Exception as e:
            logger.error(f"An unexpected error occurred while sending email to {recipient_email}: {e}")
            return False


async def send_welcome_email(recipient_email, username):
    if not config.welcome_template:
        logger.error("Welcome email template is not loaded. Cannot send welcome email.")
        return

    subject = "Welcome to Ally, your AI Meeting Wizard!"
    app_url = "https://ally-frontend-vw00.onrender.com"
    html_content = config.welcome_template.replace("[USER_NAME]", html.escape(username))
    html_content = html_content.replace("[MY_URL]", app_url)
    html_content = html_content.replace("[CURRENT_YEAR]", str(datetime.now().year))
    
    await send_email_via_brevo(recipient_email, subject, html_content, sender_name="Ally")


async def send_reminder_email(task_data, user_id, task_id, reminder_type):
    if not config.reminder_template:
        logger.error("Reminder email template is not loaded. Cannot send reminder.")
        return

    recipient_email = task_data.get("assigneeEmail")
    if not recipient_email:
        logger.warning(f"Task {task_id} has no assignee email. Skipping reminder.")
        return

    subject, reminder_message = "", ""
    base_url = os.getenv("MY_API", "https://ally-backend-y2pq.onrender.com")
    update_url = f"{base_url}/update-task-status"
    in_progress_url = f"{update_url}?userId={user_id}&taskId={task_id}&newStatus=In Progress"
    completed_url = f"{update_url}?userId={user_id}&taskId={task_id}&newStatus=Completed"
    task_title = task_data.get('title', 'N/A')
    start_date_str = task_data.get('startDate', 'N/A')
    deadline_str = task_data.get('deadline', 'N/A')

    if reminder_type == "start_date":
        subject = f"Reminder: Task '{task_title}' is scheduled to start today!"
        reminder_message = f"This is a reminder that your task, <strong>{task_title}</strong>, is scheduled to begin today, {start_date_str}."
    elif reminder_type == "before_deadline":
        subject = f"Reminder: Task '{task_title}' is due tomorrow!"
        reminder_message = f"This is a friendly reminder that your task, <strong>{task_title}</strong>, is due tomorrow, {deadline_str}."
    elif reminder_type == "deadline":
        subject = f"URGENT: Task '{task_title}' is due today!"
        reminder_message = f"This is an urgent reminder that your task, <strong>{task_title}</strong>, is due today, {deadline_str}."
    elif reminder_type == "manual_notification":
        subject = f"Notification: A message about your task '{task_title}'"
        reminder_message = f"This is a notification regarding your task: <strong>{task_title}</strong>. Please review the details and update its status as needed."

    email_body = config.reminder_template
    email_body = email_body.replace("[REMINDER_SUBJECT]", html.escape(subject))
    email_body = email_body.replace("[PREHEADER_TEXT]", html.escape(subject))
    email_body = email_body.replace("[ASSIGNEE_NAME]", html.escape(task_data.get('assignee', 'Team Member')))
    email_body = email_body.replace("[REMINDER_MESSAGE_HTML]", reminder_message)
    email_body = email_body.replace("[TASK_TITLE]", html.escape(task_title))
    email_body = email_body.replace("[TASK_START_DATE]", html.escape(start_date_str))
    email_body = email_body.replace("[TASK_DEADLINE]", html.escape(deadline_str))
    email_body = email_body.replace("[IN_PROGRESS_URL]", in_progress_url)
    email_body = email_body.replace("[COMPLETED_URL]", completed_url)
    email_body = email_body.replace("[CURRENT_YEAR]", str(datetime.now().year))
    
    await send_email_via_brevo(recipient_email, subject, email_body, sender_name="Ally, your AI Meeting Wizard")
